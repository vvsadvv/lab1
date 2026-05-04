from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "report" / "LAB2_REPORT.md"
OUTPUT_DOCX = ROOT / "report" / "LAB2_REPORT.docx"


def add_code_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def main() -> None:
    doc = Document()
    text = SOURCE_MD.read_text(encoding="utf-8")
    in_code = False

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            add_code_line(doc, line)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line.strip() == "":
            doc.add_paragraph("")
            continue
        doc.add_paragraph(line)

    doc.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    main()
