from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "report" / "REPORT.md"
OUTPUT_DOCX = ROOT / "report" / "REPORT.docx"


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    in_code_block = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_code_paragraph(doc, line)
            continue

        if stripped == "":
            doc.add_paragraph("")
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            alt_text = image_match.group(1).strip() or "image"
            image_path = image_match.group(2).strip()
            resolved = (REPORT_MD.parent / image_path).resolve()
            if resolved.exists():
                doc.add_picture(str(resolved), width=Inches(6.2))
            else:
                doc.add_paragraph(f"[Отсутствует изображение: {Path(image_path).name}]")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(text, style="List Number")
            continue

        doc.add_paragraph(stripped)


def main() -> None:
    if not REPORT_MD.exists():
        raise FileNotFoundError(f"Missing {REPORT_MD}")

    markdown_text = REPORT_MD.read_text(encoding="utf-8")
    doc = Document()
    add_markdown_to_docx(doc, markdown_text)
    doc.save(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    main()
